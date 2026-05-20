#!/usr/bin/env python3
"""Build TeamX_report.docx directly using python-docx (no pandoc required).
Saves to d:/Downloads/.github/Report/Template/docx/TeamX_report.docx
"""

import subprocess, sys, pathlib

# install deps if missing
for pkg in ['python-docx', 'lxml']:
    try:
        __import__(pkg.replace('-', '.').split('.')[0] if pkg != 'python-docx' else 'docx')
    except ImportError:
        subprocess.check_call([sys.executable, '-m', 'pip', 'install', pkg])

from docx import Document
from docx.shared import Mm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from copy import deepcopy
from lxml import etree
import zipfile, tempfile, os, shutil

OUTPUT = pathlib.Path(r'd:/Downloads/.github/Report/Template/docx/TeamX_report.docx')
OUTPUT.parent.mkdir(parents=True, exist_ok=True)

NS = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

def qname(tag):
    prefix, name = tag.split(':')
    return f'{{{NS[prefix]}}}{name}'

def ensure_rpr(style):
    rpr = style._element.rPr
    if rpr is None:
        rpr = OxmlElement('w:rPr')
        style._element.append(rpr)
    return rpr

def set_font(style, latin='Times New Roman', east_asia='Malgun Gothic', size=10, bold=None, italic=None):
    rpr = ensure_rpr(style)
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    style.font.name = latin
    rfonts.set(qn('w:ascii'), latin)
    rfonts.set(qn('w:hAnsi'), latin)
    rfonts.set(qn('w:eastAsia'), east_asia)
    if size: style.font.size = Pt(size)
    if bold is not None: style.font.bold = bold
    if italic is not None: style.font.italic = italic

# ---- build base document ----
doc = Document()

# Page setup
sec = doc.sections[0]
sec.page_width  = Mm(210)
sec.page_height = Mm(297)
sec.top_margin    = Mm(19.1)
sec.bottom_margin = Mm(25.4)
sec.left_margin   = Mm(17)
sec.right_margin  = Mm(17)

# Styles
normal = doc.styles['Normal']
set_font(normal, size=10)
normal.paragraph_format.line_spacing = Pt(13)
normal.paragraph_format.space_after  = Pt(2)

title_style = doc.styles['Title']
set_font(title_style, size=14, bold=True)
title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_style.paragraph_format.space_after = Pt(4)

h1 = doc.styles['Heading 1']
set_font(h1, size=10, bold=True)
h1.paragraph_format.space_before = Pt(10)
h1.paragraph_format.space_after  = Pt(4)

h2 = doc.styles['Heading 2']
set_font(h2, size=10, bold=True, italic=True)
h2.paragraph_format.space_before = Pt(6)
h2.paragraph_format.space_after  = Pt(3)

# ---- content ----

def add_para(doc, text, style='Normal', bold=False, italic=False, align=None, size=None, space_before=None):
    p = doc.add_paragraph(style=style)
    if align: p.alignment = align
    if space_before: p.paragraph_format.space_before = space_before
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size: run.font.size = Pt(size)
    return p

def add_rule(doc):
    """Add a thin horizontal rule paragraph."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ppr = p._element.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pbdr.append(bottom)
    ppr.append(pbdr)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    return p

# Title
p = doc.add_paragraph(style='Title')
p.add_run('2026-1 기계학습기초 팀 프로젝트 Final Report')

# Author line
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(2)
p.add_run('홍길동\u02071  김철수\u02071  이영희\u02071').font.size = Pt(10)

# Affiliation
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(2)
run = p.add_run('\u02071아주대학교 인공지능융합학과')
run.font.size = Pt(9)

# Email
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(2)
run = p.add_run('{hong, kim, lee}@ajou.ac.kr')
run.font.size = Pt(9)

# Date / supervisor
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
run = p.add_run('제출일: 2026년 6월 12일  |  담당교수: 교수명 입력')
run.font.size = Pt(9)

add_rule(doc)

# Abstract
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after  = Pt(4)
p.paragraph_format.left_indent  = Mm(6)
p.paragraph_format.right_indent = Mm(6)
run = p.add_run('Abstract.  ')
run.bold = True
run.font.size = Pt(9)
run2 = p.add_run(
    '본 보고서는 팀 프로젝트의 문제 정의, 관련 연구, 방법론, 실험 설정, 주요 결과, '
    '해석 및 한계를 종합적으로 정리한 최종 보고서이다. 프로젝트의 목적과 기여를 명확히 제시하고, '
    '재현 가능한 수준으로 실험 과정을 서술하며, 결과에 대한 분석과 향후 개선 방향까지 포함하여 '
    '작성한다. Abstract는 약 150–200 단어로 간결하게 작성한다.'
)
run2.font.size = Pt(9)

# Keywords
p = doc.add_paragraph()
p.paragraph_format.space_after  = Pt(4)
p.paragraph_format.left_indent  = Mm(6)
p.paragraph_format.right_indent = Mm(6)
run = p.add_run('Keywords: ')
run.bold = True
run.font.size = Pt(9)
run2 = p.add_run('팀 프로젝트, 기계학습, 실험 분석, 결과 해석, Git 협업')
run2.font.size = Pt(9)

add_rule(doc)

SECTIONS = [
    ('1. INTRODUCTION',
     '이 절에서는 프로젝트가 다루는 문제의 배경과 필요성을 설명한다. 왜 이 문제가 중요한지, '
     '어떤 실제적 또는 학술적 의미가 있는지를 서술하고, 본 프로젝트에서 해결하고자 하는 과제를 '
     '명확히 정의한다. 또한 프로젝트의 주요 목표와 핵심 기여를 요약한다.\n\n'
     '포함 권장 흐름: 문제 배경과 동기 / 해결하고자 하는 예측 또는 분류 과제 정의 / '
     '프로젝트의 주요 목표와 기여 요약 / 이하 절의 구성 안내'),
    ('2. RELATED WORK',
     '이 절에서는 기존 연구, 유사한 접근 방식, 혹은 이 문제를 이해하기 위해 필요한 배경지식을 '
     '정리한다. 학술 논문, 공개 프로젝트, benchmark, baseline 모델 등을 조사하여 본 프로젝트가 '
     '어떤 맥락에 위치하는지 설명한다. 필요에 따라 이 절의 제목은 Background 또는 '
     'Problem Setting 등으로 조정할 수 있다.'),
    ('3. APPROACH',
     '이 절에서는 문제를 해결하기 위해 사용한 전체 접근 방법을 설명한다. 전처리 전략, 모델 선택, '
     '학습 절차, 하이퍼파라미터 설정 등을 독자가 이해할 수 있도록 서술한다.\n\n'
     '포함 권장 항목: 전체 문제 해결 전략 및 파이프라인 개요 / 데이터 전처리 방법 / '
     '사용 모델 및 알고리즘 설명 / 하이퍼파라미터 설정 근거 / 사용한 라이브러리 및 구현 환경'),
]

for heading, body in SECTIONS:
    doc.add_heading(heading, level=1)
    for para_text in body.split('\n\n'):
        p = doc.add_paragraph(para_text.strip())
        p.paragraph_format.space_after = Pt(3)

# Section 4: Experiment with tables
doc.add_heading('4. EXPERIMENT', level=1)
p = doc.add_paragraph(
    '이 절에서는 실험 설계와 사용한 데이터셋을 구체적으로 설명한다. '
    '동일한 조건에서 여러 모델을 비교할 수 있도록 데이터 분할, 실험 환경, 평가 지표를 명시한다.'
)
p.paragraph_format.space_after = Pt(3)

doc.add_heading('4.1 Dataset', level=2)

tbl1 = doc.add_table(rows=6, cols=2)
tbl1.style = 'Table Grid'
headers1 = ['항목', '내용']
rows1 = [
    ('데이터셋', '예: UCI Adult / Kaggle 공개 데이터'),
    ('예측 목표', '예: 소득 수준 분류'),
    ('표본 수',   '예: 10,000'),
    ('피처 수',   '예: 18'),
    ('출처',      'URL, 날짜, 라이선스 명시'),
]
for i, h in enumerate(headers1):
    cell = tbl1.rows[0].cells[i]
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True
for r, (c0, c1) in enumerate(rows1, start=1):
    tbl1.rows[r].cells[0].text = c0
    tbl1.rows[r].cells[1].text = c1

doc.add_paragraph()

doc.add_heading('4.2 Experimental Setup', level=2)
tbl2 = doc.add_table(rows=6, cols=2)
tbl2.style = 'Table Grid'
rows2 = [
    ('항목',      '설정'),
    ('데이터 분할', 'train 70% / val 15% / test 15%'),
    ('평가지표',   'Accuracy, Macro-F1'),
    ('베이스라인',  'Logistic Regression'),
    ('비교 모델',  'Random Forest, XGBoost'),
    ('실험 환경',  'Python 3.10, scikit-learn 1.4'),
]
for r, (c0, c1) in enumerate(rows2):
    tbl2.rows[r].cells[0].text = c0
    tbl2.rows[r].cells[1].text = c1
    if r == 0:
        for cell in tbl2.rows[r].cells:
            cell.paragraphs[0].runs[0].bold = True

doc.add_paragraph()

# Section 5: Results
doc.add_heading('5. RESULTS', level=1)
p = doc.add_paragraph(
    '이 절에서는 주요 정량적 결과를 표와 그림 중심으로 제시한다. '
    '어떤 모델이 어떤 지표에서 우수했는지 명확하게 드러나도록 구성한다.'
)
p.paragraph_format.space_after = Pt(3)

tbl3 = doc.add_table(rows=4, cols=4)
tbl3.style = 'Table Grid'
rows3 = [
    ('모델', 'Accuracy', 'F1-score', '비고'),
    ('Baseline (LR)', '0.812', '0.801', '기본 설정'),
    ('Random Forest', '0.846', '0.838', 'feat. eng. 적용'),
    ('XGBoost',       '0.861', '0.854', '최종 선택 모델'),
]
for r, row_data in enumerate(rows3):
    for c, val in enumerate(row_data):
        cell = tbl3.rows[r].cells[c]
        cell.text = val
        if r == 0:
            cell.paragraphs[0].runs[0].bold = True

doc.add_paragraph()

# Sections 6 & 7
doc.add_heading('6. ANALYSIS / DISCUSSION', level=1)
p = doc.add_paragraph(
    '이 절에서는 결과를 해석하고, 왜 그런 결과가 나왔는지 분석한다. 성능이 잘 나온 이유와 '
    '잘 나오지 않은 이유를 모두 다루며, 실패 사례와 한계점도 함께 정리한다. '
    '단순히 "성능이 높았다/낮았다"로 끝내지 말고, 가능한 근거를 들어 설명한다.\n\n'
    '포함 권장 항목: 성능 차이에 대한 원인 분석 / 중요한 feature 또는 모델 특성 해석 / '
    '오분류/오예측 사례 분석 / 데이터 품질 또는 실험 설계의 한계 / 개선 방향 및 후속 연구 아이디어'
)
p.paragraph_format.space_after = Pt(3)

doc.add_heading('7. CONCLUSION', level=1)
p = doc.add_paragraph(
    '이 절에서는 프로젝트 전체를 간단히 요약한다. 어떤 문제를 다뤘고, 어떤 방법을 사용했으며, '
    '어떤 결과와 의미를 얻었는지를 정리한다. 마지막으로 향후 확장 가능성이나 실제 적용 가능성을 '
    '짧게 덧붙일 수 있다.'
)

doc.add_heading('References', level=1)
p = doc.add_paragraph(
    '참고문헌은 references.bib를 사용하여 정리한다. 학술 논문, 데이터셋, 오픈소스 라이브러리, '
    '공식 문서 등을 일관된 형식으로 인용한다.\n\n'
    '[1] L. Breiman, "Random Forests," Machine Learning, vol. 45, no. 1, pp. 5–32, 2001.\n'
    '[2] T. Chen and C. Guestrin, "XGBoost: A Scalable Tree Boosting System," KDD 2016.\n'
    '[3] F. Pedregosa et al., "Scikit-learn: Machine Learning in Python," JMLR, vol. 12, 2011.\n'
    '[4] D. Dua and C. Graff, UCI Machine Learning Repository, 2019.'
)
p.paragraph_format.space_after = Pt(3)

doc.add_heading('Appendix (Optional)', level=1)
p = doc.add_paragraph(
    '필요한 경우 부록을 추가할 수 있다. 예시: 추가 실험 결과 / 하이퍼파라미터 상세표 / '
    '데이터 전처리 세부 절차 / 모델 구조 그림 / 역할 분담 및 GitHub 협업 로그 요약'
)

# ---- post-process: insert single→two-column section break ----------------
raw_path = str(OUTPUT).replace('.docx', '.raw.docx')
doc.save(raw_path)

with tempfile.TemporaryDirectory() as tmp:
    with zipfile.ZipFile(raw_path) as zin:
        zin.extractall(tmp)
    
    doc_xml = os.path.join(tmp, 'word', 'document.xml')
    root = etree.parse(doc_xml).getroot()
    body = root.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}body')
    final_sect = body.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr')
    
    if final_sect is not None:
        # Find first heading paragraph
        first_h = None
        for para in body.findall('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p'):
            ps = para.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pStyle')
            if ps is not None and 'Heading' in ps.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', ''):
                first_h = para
                break
        
        if first_h is not None:
            W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            # Section break paragraph (single-col header ends here)
            sb_para = etree.Element(f'{{{W}}}p')
            sb_ppr  = etree.SubElement(sb_para, f'{{{W}}}pPr')
            sb_sect = etree.SubElement(sb_ppr,  f'{{{W}}}sectPr')
            for child_tag in [f'{{{W}}}pgSz', f'{{{W}}}pgMar']:
                child = final_sect.find(child_tag)
                if child is not None:
                    sb_sect.append(deepcopy(child))
            # single column (no cols element = 1 col)
            body.insert(list(body).index(first_h), sb_para)
            
            # Make body two-column
            cols = final_sect.find(f'{{{W}}}cols')
            if cols is None:
                cols = etree.SubElement(final_sect, f'{{{W}}}cols')
            cols.attrib.clear()
            cols.set(f'{{{W}}}num', '2')
            cols.set(f'{{{W}}}space', '510')
    
    with open(doc_xml, 'wb') as f:
        f.write(etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone=True))
    
    with zipfile.ZipFile(str(OUTPUT), 'w', zipfile.ZIP_DEFLATED) as zout:
        for folder, _, files in os.walk(tmp):
            for file in files:
                full = os.path.join(folder, file)
                rel  = os.path.relpath(full, tmp)
                zout.write(full, rel)

os.remove(raw_path)
print(f'Done: {OUTPUT}')
